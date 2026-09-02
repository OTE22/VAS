"""Rendering documents, and persisting them as artifacts.

Moved out of sql_agent/api/routes.py so the agent's graph nodes can render a
document without importing the API layer (routes imports the agent, so the
other direction would be a cycle). The rendering code itself is unchanged —
same reportlab/python-docx calls, same sanitisation, same Arabic shaping — it
only returns BYTES now instead of an HTTP Response, because bytes are what
both callers need: the endpoint wraps them in a Response, the graph node
hands them to the registry.

`render_and_register` is the single persistence path. There is deliberately
not a second one: an artifact created by the HTTP export and an artifact
created by the agent are the same kind of object, with the same ownership,
retention and lineage guarantees, so they are created by the same function.
"""

import asyncio
import logging
import os
import re
from datetime import datetime
from io import BytesIO
from typing import Optional, Tuple

from fastapi import HTTPException
from pydantic import BaseModel
from sqlalchemy.exc import IntegrityError

from sql_agent.services import artifact_registry

logger = logging.getLogger(__name__)


# Export Models — server-enforced size limits (Pydantic rejects oversize)
class ExportRequest(BaseModel):
    content: str  # validated in _sanitize_export below
    title: str
    timestamp: str


_EXPORT_MAX_CONTENT_CHARS = 500_000
_EXPORT_MAX_TITLE_CHARS = 200
# Bounded concurrent document generation (reportlab/docx are CPU-bound)
_export_semaphore = asyncio.Semaphore(4)


def sanitize_export(request: "ExportRequest") -> Tuple[str, str, str]:
    """Size limits + markup-injection prevention for document generation.

    reportlab's Paragraph parses XML-ish markup — raw '<'/'&' from the model
    or browser must be escaped BEFORE we selectively re-allow <b>/<i>.
    Returns (safe_title, safe_content, safe_date) or raises 413/422.
    """
    from xml.sax.saxutils import escape as _xml_escape
    if len(request.content) > _EXPORT_MAX_CONTENT_CHARS:
        raise HTTPException(status_code=413, detail="Export content too large")
    title = re.sub(r'[<>&\x00-\x1f]', '', request.title or 'Intelligence Report').strip()
    title = title[:_EXPORT_MAX_TITLE_CHARS] or 'Intelligence Report'
    content = _xml_escape(request.content)
    # Safe date for the filename (never trust raw client timestamp strings)
    safe_date = datetime.utcnow().strftime('%Y-%m-%d')
    return title, content, safe_date

_ARABIC_CHAR = re.compile(r"[؀-ۿ]")
_INLINE_TAG = re.compile(r"</?(?:b|i)>")


def _register_pdf_font():
    """Register the repo-shipped Arabic-capable font with reportlab, once.

    Returns the font name to use, or None to stay on Helvetica. The TTF is
    merged from the SAME vendored Cairo subsets the web UI renders with
    (assets/fonts/Cairo-PDF.ttf), so PDF and screen agree — and its cmap maps
    the Arabic presentation forms the reshaper emits, which is the property
    that makes shaped text render instead of .notdef boxes.
    """
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont as _RLTTFont
        if "CairoPDF" in pdfmetrics.getRegisteredFontNames():
            return "CairoPDF"
        font_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
            "assets", "fonts", "Cairo-PDF.ttf")
        if not os.path.isfile(font_path):
            return None
        pdfmetrics.registerFont(_RLTTFont("CairoPDF", font_path))
        return "CairoPDF"
    except Exception:
        logger.warning("[EXPORT] Arabic-capable PDF font unavailable; "
                       "falling back to Helvetica", exc_info=True)
        return None


def _shape_rtl(text: str) -> str:
    """Shape + bidi-reorder one LINE of Arabic-bearing text for reportlab.

    reportlab draws glyphs in the order given with no complex text layout, so
    Arabic needs its positional letter forms substituted (arabic_reshaper) and
    the line reordered for display (python-bidi). Inline <b>/<i> markers are
    stripped first: bidi reordering would scatter the ASCII tags through the
    RTL text, which is worse than losing bold.

    Fail-open: if the shaping libraries are missing, the raw text is returned
    — wrong-looking Arabic beats a failed export.
    """
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        return get_display(arabic_reshaper.reshape(_INLINE_TAG.sub("", text)))
    except Exception:
        return text

#: Bold, then italic. Bold first so ** is consumed before a single * can
#: match half of it. Both require a non-empty body on ONE line, so a stray
#: asterisk ("5 * 3") stays literal instead of swallowing the rest.
_MD_BOLD = re.compile(r"\*\*(?=\S)(.+?)(?<=\S)\*\*")
_MD_ITALIC = re.compile(r"(?<!\*)\*(?=\S)([^*\n]+?)(?<=\S)\*(?!\*)")
_MD_CODE = re.compile(r"`([^`\n]+)`")
_MD_HEADING = re.compile(r"^(#{1,6})\s*(.*)$")


def _markdown_inline(text: str) -> str:
    """Inline marks to reportlab tags.

    Operates on ALREADY-ESCAPED text: the caller escapes '<' and '&' because
    reportlab parses XML-ish markup, and that must not be undone here. Only
    the <b>/<i> this function introduces are real tags.
    """
    text = _MD_BOLD.sub(r"<b>\1</b>", text)
    text = _MD_ITALIC.sub(r"<i>\1</i>", text)
    return _MD_CODE.sub(r"\1", text)


def _markdown_block(line: str) -> tuple:
    """(heading level, text) for one line, with the marker REMOVED.

    Level 0 is body text. Quotes and list items keep their content and lose
    their marker - a printed '>' or '-' is exactly what made the report look
    unfinished.
    """
    raw = (line or "").strip()

    heading = _MD_HEADING.match(raw)
    if heading:
        return len(heading.group(1)), _markdown_inline(heading.group(2).strip())

    if raw.startswith(">"):
        return 0, _markdown_inline(raw.lstrip(">").strip())

    if raw[:2] in ("- ", "* ") or raw[:2] == "\u2022 ":
        return 0, "\u2022 " + _markdown_inline(raw[2:].strip())

    return 0, _markdown_inline(raw)


def build_pdf_bytes(safe_title: str, safe_content: str, safe_date: str, analyst: str):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        # Create PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, 
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=72)
        
        # Container for PDF content
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#00ff96'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#00ff96'),
            spaceAfter=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.black,
            leading=14,
            spaceAfter=12
        )

        # Arabic support. Helvetica has no Arabic glyphs — an Arabic report
        # exported as nothing but ■■■■ boxes. When the merged Cairo font is
        # available, Arabic-bearing paragraphs are shaped, bidi-reordered and
        # set right-aligned in it; Latin paragraphs keep the original look.
        arabic_font = _register_pdf_font()
        from reportlab.lib.enums import TA_RIGHT
        rtl_body_style = ParagraphStyle(
            'RTLBody', parent=body_style,
            fontName=arabic_font or body_style.fontName,
            alignment=TA_RIGHT, wordWrap='RTL',
        )
        rtl_header_style = ParagraphStyle(
            'RTLHeader', parent=header_style,
            fontName=arabic_font or header_style.fontName,
            alignment=TA_RIGHT, wordWrap='RTL',
        )

        # Add title
        title = Paragraph(f"<b>INTELLIGENCE REPORT</b>", title_style)
        story.append(title)
        story.append(Spacer(1, 0.2*inch))

        # Add metadata (all values pre-sanitized in _sanitize_export).
        # The query itself may be Arabic (e.g. "تتبع IRON MAN") — shape just
        # that value; the labels stay Latin.
        shown_title = safe_title
        if arabic_font and _ARABIC_CHAR.search(safe_title):
            shown_title = _shape_rtl(safe_title)
        meta_style = (rtl_body_style if arabic_font and _ARABIC_CHAR.search(safe_title)
                      else body_style)
        metadata = f"<b>Query:</b> {shown_title}<br/>"
        metadata += f"<b>Generated:</b> {safe_date}<br/>"
        metadata += f"<b>Analyst:</b> {re.sub(r'[<>&]', '', analyst)}"
        story.append(Paragraph(metadata, meta_style))
        story.append(Spacer(1, 0.3*inch))

        # Content arrives XML-escaped; re-allow ONLY escaped bold/italic markers
        content = safe_content
        content = content.replace('&lt;br&gt;', '\n').replace('&lt;br/&gt;', '\n')
        content = re.sub(r'&lt;strong&gt;(.*?)&lt;/strong&gt;', r'<b>\1</b>', content)
        content = re.sub(r'&lt;em&gt;(.*?)&lt;/em&gt;', r'<i>\1</i>', content)

        # Split into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if not para.strip():
                continue
            raw = para.strip()
            # Convert the MARKUP before anything else: the marker decides the
            # style, and the marker itself must not reach the page.
            blocks = [_markdown_block(line) for line in raw.split('\n')]
            level = blocks[0][0] if blocks else 0

            if arabic_font and _ARABIC_CHAR.search(raw):
                # Shaping is PER LINE: bidi reordering operates on a line, and
                # a timestamp at line start must stay attached to its own
                # line, not migrate across breaks.
                text = '<br/>'.join(_shape_rtl(body) for _lvl, body in blocks)
                story.append(Paragraph(text, rtl_header_style if level
                                       else rtl_body_style))
            else:
                text = '<br/>'.join(body for _lvl, body in blocks)
                if level == 1:
                    style = header_style
                elif level >= 2:
                    style = header_style
                else:
                    style = body_style
                story.append(Paragraph(text, style))
            story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF bytes
        buffer.seek(0)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        return pdf_bytes
    except ImportError:
        logger.error("[EXPORT] reportlab not installed. Install with: pip install reportlab")
        raise HTTPException(
            status_code=500,
            detail="PDF export is not available on this server."
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[EXPORT] PDF export error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate PDF")

def build_word_bytes(safe_title: str, raw_content: str, safe_date: str, analyst: str):
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create Word document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('INTELLIGENCE REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(0, 255, 150)  # #00ff96
        
        # Add metadata (sanitized title, server date; docx runs are plain text)
        doc.add_paragraph(f'Query: {safe_title}')
        doc.add_paragraph(f'Generated: {safe_date}')
        doc.add_paragraph(f'Analyst: {analyst}')
        doc.add_paragraph()  # Empty line

        # Add content
        content = raw_content
        # Clean HTML
        content = re.sub(r'<br\s*/?>', '\n', content)
        content = re.sub(r'<strong>(.*?)</strong>', r'\1', content)  # Remove strong tags
        content = re.sub(r'<em>(.*?)</em>', r'\1', content)  # Remove em tags
        content = re.sub(r'<[^>]+>', '', content)  # Remove remaining HTML tags
        
        # Split into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.style.font.size = Pt(11)
                p.style.font.name = 'Calibri'
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        word_bytes = buffer.getvalue()
        buffer.close()
        
        return word_bytes
    except ImportError:
        logger.error("[EXPORT] python-docx not installed. Install with: pip install python-docx")
        raise HTTPException(
            status_code=500,
            detail="Word export is not available on this server."
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[EXPORT] Word export error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate Word document")

# Names the API layer has always used, kept so its call sites stay unchanged.
_build_pdf_export = build_pdf_bytes
_build_word_export = build_word_bytes
_sanitize_export = sanitize_export


# ---------------------------------------------------------------- persistence

async def render_and_register(
    db, *, payload: bytes, artifact_type: str, title: str, language: str,
    user_id, created_by_username: str, conversation_id=None,
    source_query: Optional[str] = None, source_sql: Optional[str] = None,
    source_content: Optional[str] = None, source_message_id=None,
    source_result_id: Optional[int] = None, modification_meta: Optional[dict] = None,
    parent_artifact_id=None,
) -> Optional[str]:
    """Persist rendered bytes as an artifact. Returns its id, or None.

    Returning None rather than raising is deliberate, and it is the whole
    contract of the HTTP export path: the caller already holds a finished
    document, and refusing to hand it over because a bookkeeping row failed
    would turn a working feature into an outage. So a failure here costs the
    user the ability to say "translate the last report" — it never costs them
    the report.

    What it must NOT do is half-succeed. register_artifact writes the file
    first, then the row, and unlinks the file (and expunges the pending row
    from the caller's session) if the row fails — so None means nothing was
    persisted, not "something was".

    A STALE LINEAGE POINTER COSTS THE LINEAGE, NEVER THE ARTIFACT. Working
    memory durably remembers the history row a result came from; retention,
    user deletion or a test teardown can remove that row while the pointer
    lives on in the session file. Inserting then dies on the foreign key —
    and because the stale pointer persists, every LATER document would fail
    the same way. On an IntegrityError the insert is retried exactly once
    with the optional lineage references stripped: the user gets their
    document and a registered artifact; only the dangling provenance link is
    lost, which is the truthful outcome — the row it named no longer exists.
    """
    async def _attempt(message_id, result_id, parent_id):
        artifact = await artifact_registry.register_artifact(
            db, payload=payload, artifact_type=artifact_type, title=title,
            language=language, user_id=user_id,
            created_by_username=created_by_username,
            conversation_id=conversation_id, source_query=source_query,
            source_sql=source_sql, source_content=source_content,
            source_message_id=message_id, source_result_id=result_id,
            modification_meta=modification_meta, parent_artifact_id=parent_id,
        )
        await db.commit()
        return str(artifact.id)

    try:
        return await _attempt(source_message_id, source_result_id,
                              parent_artifact_id)
    except IntegrityError as e:
        try:
            await db.rollback()
        except Exception:
            pass
        if not (source_message_id or source_result_id or parent_artifact_id):
            # Nothing optional to strip: the violation is something else
            # (user_id, conversation_id) and retrying would just fail again.
            logger.warning("[EXPORT] artifact not registered "
                           "(document still served): %s", e)
            return None
        logger.warning(
            "[EXPORT] lineage reference no longer exists "
            "(message_id=%s result_id=%s parent=%s); registering the artifact "
            "without it: %s",
            source_message_id, source_result_id, parent_artifact_id, e)
        try:
            return await _attempt(None, None, None)
        except Exception as retry_error:
            logger.warning("[EXPORT] artifact not registered "
                           "(document still served): %s", retry_error)
            try:
                await db.rollback()
            except Exception:
                pass
            return None
    except Exception as e:
        logger.warning("[EXPORT] artifact not registered (document still served): %s", e)
        try:
            await db.rollback()
        except Exception:
            pass
        return None


def detect_language(text: str) -> str:
    """'ar' if the document contains Arabic, else 'en'.

    The same test the PDF renderer uses to decide on shaping, so a document
    that RENDERS as Arabic is also RECORDED as Arabic — a translation request
    against it then has the right starting point.
    """
    return "ar" if _ARABIC_CHAR.search(text or "") else "en"
