"""
SQL Agent API Routes
====================
FastAPI router for SQL Intelligence Agent endpoints.
"""

import asyncio
import json
import logging
import threading
from collections import OrderedDict
from datetime import datetime
from typing import Optional, Dict

from fastapi import APIRouter, WebSocket, WebSocketDisconnect, HTTPException, Depends, Request
from fastapi.responses import JSONResponse, StreamingResponse, Response
from fastapi.concurrency import run_in_threadpool
from pydantic import BaseModel
from io import BytesIO
import re

# Add parent directory to path for auth imports
import os
import sys
parent_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
if parent_dir not in sys.path:
    sys.path.insert(0, parent_dir)

try:
    from backend.auth.auth_service import get_current_user, require_chatbot_access
    from db_models import User
    from db_connection import get_db
    from sql_agent.services.user_query_history_service import user_query_history_service
    AUTH_AVAILABLE = True
except ImportError:
    AUTH_AVAILABLE = False
    # Create dummy dependencies if auth is not available
    def get_current_user():
        return None
    def require_chatbot_access():
        def dummy_dep():
            return None
        return dummy_dep

logger = logging.getLogger(__name__)

# Router for SQL Agent endpoints
router = APIRouter(prefix="/api/sql-agent", tags=["SQL Agent"])

# Global state - will be set by backend.lifespan during startup
_sql_agent_instance = None
_sql_agent_available = False

# ---------------------------------------------------------------------------
# SQL-agent isolation: concurrency cap, total timeout, per-user serialization,
# bounded LRU of user agents, and a thread->asyncio.Queue streaming bridge.
# The LangGraph pipeline is fully synchronous (sync ChatOllama + psycopg2), so
# every query MUST run in its own thread — never on the event loop — and the
# number of simultaneous queries must be small (each one monopolizes Ollama).
# ---------------------------------------------------------------------------
try:
    from config import settings as _app_settings
    SQL_AGENT_MAX_CONCURRENT = int(getattr(_app_settings, 'SQL_AGENT_MAX_CONCURRENT', 2))
    SQL_AGENT_TOTAL_TIMEOUT = float(getattr(_app_settings, 'SQL_AGENT_TOTAL_TIMEOUT', 300))
except Exception:
    SQL_AGENT_MAX_CONCURRENT = 2
    SQL_AGENT_TOTAL_TIMEOUT = 300.0

_sql_agent_semaphore = asyncio.Semaphore(SQL_AGENT_MAX_CONCURRENT)
_SEMAPHORE_WAIT_SECONDS = 5.0
_BUSY_MESSAGE = (
    "The SQL assistant is currently handling other requests. "
    "Please try again in a few moments."
)

# User-specific agent instances (one per user for persistent memory).
# Bounded LRU: at most _USER_AGENTS_MAX users keep a live agent; the least
# recently used is evicted (previously this dict grew forever).
_USER_AGENTS_MAX = 10
_user_agents: "OrderedDict[int, object]" = OrderedDict()

# Per-user locks so two concurrent queries from the SAME user (double-submit,
# two tabs) serialize instead of corrupting shared conversation memory.
_user_query_locks: Dict[int, asyncio.Lock] = {}


def _get_user_lock(user_id: int) -> asyncio.Lock:
    lock = _user_query_locks.get(user_id)
    if lock is None:
        lock = asyncio.Lock()
        _user_query_locks[user_id] = lock
    return lock


def _get_or_create_user_agent(user_id: int):
    """Get the user's agent (LRU-refresh) or create it, evicting the oldest."""
    agent = _user_agents.get(user_id)
    if agent is not None:
        _user_agents.move_to_end(user_id)
        return agent

    from sql_agent.agent import SQLIntelligenceAgent
    from sql_agent.conversation_memory import ConversationMemory

    logger.info(f"[SQL_AGENT_API] Creating user-specific agent for user {user_id}")
    user_memory = ConversationMemory(user_id=user_id)
    user_session_id = user_memory.start_session()
    agent = SQLIntelligenceAgent(conversation_memory=user_memory)
    _user_agents[user_id] = agent
    logger.info(f"[SQL_AGENT_API] User {user_id} agent created (session: {user_session_id})")

    while len(_user_agents) > _USER_AGENTS_MAX:
        evicted_id, _evicted = _user_agents.popitem(last=False)
        _user_query_locks.pop(evicted_id, None)
        logger.info(f"[SQL_AGENT_API] Evicted LRU agent for user {evicted_id} (cap {_USER_AGENTS_MAX})")

    return agent


_STREAM_SENTINEL = object()

# ---------------------------------------------------------------------------
# Request registry (correlation, cancellation, idempotency)
# ---------------------------------------------------------------------------
# request_id -> {"cancel_event": threading.Event, "user_id": int, "status": str,
#                "started_at": float}
# Bounded: terminal entries are pruned; hard cap evicts oldest.
_ACTIVE_REQUESTS: "OrderedDict[str, dict]" = OrderedDict()
_MAX_TRACKED_REQUESTS = 300
_REQUEST_ID_RE = re.compile(r'^[A-Za-z0-9_-]{8,64}$')


def _normalize_request_id(raw) -> str:
    """Accept a client-supplied request id (uuid-ish) or mint one."""
    import uuid as _uuid
    if isinstance(raw, str) and _REQUEST_ID_RE.match(raw):
        return raw
    return _uuid.uuid4().hex


def _register_request(request_id: str, user_id, cancel_event: threading.Event) -> bool:
    """Register an accepted request. Returns False if this request_id was
    already accepted (idempotency: a fallback transport must NOT re-execute)."""
    existing = _ACTIVE_REQUESTS.get(request_id)
    if existing is not None:
        return False
    while len(_ACTIVE_REQUESTS) >= _MAX_TRACKED_REQUESTS:
        _ACTIVE_REQUESTS.popitem(last=False)
    _ACTIVE_REQUESTS[request_id] = {
        "cancel_event": cancel_event,
        "user_id": user_id,
        "status": "running",
        "started_at": asyncio.get_event_loop().time(),
    }
    return True


def _finish_request(request_id: str, status: str):
    entry = _ACTIVE_REQUESTS.get(request_id)
    if entry is not None:
        entry["status"] = status  # kept (bounded) so duplicate submits still 409


def _sse_event(payload: dict, request_id: str, seq: int) -> str:
    """Every server event repeats request_id + sequence (client correlation)."""
    payload = {**payload, "request_id": request_id, "sequence": seq}
    return f"data: {json.dumps(payload)}\n\n"


def _error_body(code: str, message: str, reference_id: str = None,
                retryable: bool = False, status: str = None) -> dict:
    """Structured error contract — the frontend reacts to `error.code`,
    NEVER to substring matching. No internal details are exposed."""
    body = {
        "success": False,
        "error": {
            "code": code,
            "message": message,
            "reference_id": reference_id,
            "retryable": retryable,
        },
        "response": None,
    }
    if status:
        body["status"] = status
    return body


def require_sql_agent_csrf(request: Request):
    """CSRF defense-in-depth for cookie-authenticated mutating requests
    (same policy as the live-alerts routes): SameSite=lax cookie + custom
    header. Bearer-token clients are exempt (token can't be sent cross-site)."""
    if request.headers.get("authorization"):
        return
    if request.headers.get("x-requested-with", "").lower() != "xmlhttprequest":
        raise HTTPException(
            status_code=403,
            detail="CSRF check failed: X-Requested-With header required",
        )


# ---------------------------------------------------------------------------
# Security policy: a denied query is Denied -> Audited -> Explained safely.
# Account blocking requires an EXPLICIT policy threshold — never a single
# model hallucination, and NEVER substring-matching the model's prose.
# ---------------------------------------------------------------------------
_SECURITY_VIOLATION_THRESHOLD = 3          # explicit denials before blocking
_SECURITY_VIOLATION_WINDOW_SECONDS = 3600.0
_security_violations: Dict[int, list] = {}  # user_id -> [monotonic timestamps]


def _record_security_violation(user_id: int) -> int:
    """Record one explicit denied operation; returns violations in window."""
    now = asyncio.get_event_loop().time()
    times = [t for t in _security_violations.get(user_id, [])
             if now - t < _SECURITY_VIOLATION_WINDOW_SECONDS]
    times.append(now)
    _security_violations[user_id] = times[-20:]
    return len(times)


async def _handle_security_denial(current_user, query: str, reason: str,
                                  execution_time_ms: float, session_id=None) -> dict:
    """Structured outcome for an explicitly denied operation.

    QUERY_DENIED for isolated denials (audited, safe explanation);
    ACCOUNT_BLOCKED only after the policy threshold is crossed.
    The rejected SQL is written to the AUDIT LOG, never to the client.
    """
    import uuid as _uuid
    reference_id = f"SEC-{_uuid.uuid4().hex[:8]}"

    if not (AUTH_AVAILABLE and current_user):
        return _error_body("QUERY_DENIED",
                           "This operation is not permitted. The assistant is read-only.",
                           reference_id)

    count = _record_security_violation(current_user.id)
    logger.warning(
        "[SECURITY] outcome=QUERY_DENIED user_id=%s violations_in_window=%d reference_id=%s",
        current_user.id, count, reference_id)

    try:
        await log_chatbot_query(
            user_id=current_user.id,
            username=current_user.username,
            query=query,
            response=None,
            success=False,
            error_message=f"SECURITY_WARNING [{reference_id}]: denied operation - {reason[:300]}",
            processing_time_ms=execution_time_ms,
            session_id=session_id,
        )
    except Exception as audit_error:
        logger.error(f"[SECURITY] audit write failed: {audit_error}")

    if count >= _SECURITY_VIOLATION_THRESHOLD:
        logger.error(
            "[SECURITY] outcome=ACCOUNT_BLOCKED user_id=%s threshold=%d reference_id=%s",
            current_user.id, _SECURITY_VIOLATION_THRESHOLD, reference_id)
        try:
            await block_user_for_forbidden_sql(
                user_id=current_user.id,
                username=current_user.username,
                sql_query=query,
                reason=f"{_SECURITY_VIOLATION_THRESHOLD} denied operations within "
                       f"{int(_SECURITY_VIOLATION_WINDOW_SECONDS / 60)} minutes [{reference_id}]",
            )
        except Exception as block_error:
            logger.error(f"[SECURITY] blocking failed: {block_error}")
        return _error_body(
            "ACCOUNT_BLOCKED",
            "Your access to the SQL assistant is temporarily restricted. "
            "Please contact an administrator.",
            reference_id)

    return _error_body(
        "QUERY_DENIED",
        "That operation is not permitted — the assistant can only read data. "
        "Repeated attempts will restrict your access.",
        reference_id)


def _start_stream_thread(agent_instance, query: str, cancel_event: threading.Event, loop) -> asyncio.Queue:
    """TRUE streaming bridge: a dedicated thread pumps the agent's blocking
    query_stream() generator into an asyncio.Queue item by item, so SSE/WS
    consumers can forward each update the moment it is produced (the previous
    implementation either collected the full list before the first byte, or —
    worse — iterated the blocking generator directly on the event loop)."""
    q: asyncio.Queue = asyncio.Queue()

    def _pump():
        try:
            for update in agent_instance.query_stream(query, cancel_event=cancel_event):
                loop.call_soon_threadsafe(q.put_nowait, update)
                if cancel_event.is_set():
                    break
        except Exception as e:
            logger.error(f"[SQL_AGENT_API] Stream thread error: {e}", exc_info=True)
            try:
                loop.call_soon_threadsafe(q.put_nowait, {"type": "error", "message": str(e)})
            except RuntimeError:
                pass  # loop closed
        finally:
            try:
                loop.call_soon_threadsafe(q.put_nowait, _STREAM_SENTINEL)
            except RuntimeError:
                pass

    threading.Thread(target=_pump, name="sql-agent-stream", daemon=True).start()
    return q

# Holds strong references to fire-and-forget background tasks. The event loop only
# keeps weak references, so without this a task can be garbage-collected mid-run.
_background_tasks: set = set()


def _spawn_background(coro):
    """Schedule a best-effort background task, keeping a strong reference to it."""
    task = asyncio.create_task(coro)
    _background_tasks.add(task)
    task.add_done_callback(_background_tasks.discard)
    return task


async def persist_query_history(
    user_id: int,
    query: str,
    response: Optional[str],
    session_id: Optional[str],
    success: bool = True,
    processing_time_ms: Optional[float] = None,
    metadata: Optional[dict] = None,
):
    """Reliably persist a query so it appears in the user's sidebar history.

    The core history row (session + query) is committed synchronously so it always
    lands. Semantic-search extras (memories + embeddings) run afterward as a
    referenced background task and can never block or fail the core save.

    Callers should ``await`` this. It is meant to run after any streamed response is
    already sent, so awaiting adds no user-facing latency while guaranteeing the row
    is committed before the request tears down (the previous fire-and-forget approach
    raced teardown and dropped queries intermittently).
    """
    if not (AUTH_AVAILABLE and user_id):
        return None

    query_history_id = None
    try:
        async for db in get_db():
            db_session = await user_query_history_service.get_or_create_session(
                db=db, user_id=user_id, session_id=session_id
            )
            query_history = await user_query_history_service.save_query_history(
                db=db,
                user_id=user_id,
                query_text=query,
                response_text=response,
                session_id=session_id,
                success=success,
                processing_time_ms=processing_time_ms,
                metadata=metadata or {},
            )
            await db.commit()
            query_history_id = query_history.id
            logger.info(f"[HISTORY_SAVE] ✅ Query {query_history_id} saved for user {user_id} (session {session_id})")
            break
    except Exception as e:
        logger.error(f"[HISTORY_SAVE] ❌ Failed to persist query history for user {user_id}: {e}", exc_info=True)
        return None

    # Best-effort enrichment (memories + embeddings) — never blocks the sidebar row
    if query_history_id is not None:
        _spawn_background(_enrich_query_history(
            user_id=user_id,
            query_history_id=query_history_id,
            query=query,
            response=response,
            session_id=session_id,
        ))
    return query_history_id


async def _enrich_query_history(user_id, query_history_id, query, response, session_id):
    """Best-effort semantic-search enrichment: memory extraction + query embedding.

    Runs after the core history row is already committed, so any failure here leaves
    the sidebar entry intact.
    """
    try:
        async for db in get_db():
            try:
                await user_query_history_service.extract_and_save_memories(
                    db=db, user_id=user_id, query_id=query_history_id,
                    query_text=query, response_text=response, session_id=session_id,
                )
            except Exception as mem_err:
                logger.warning(f"[HISTORY_ENRICH] memory extraction failed for query {query_history_id}: {mem_err}")
            try:
                embedding = await user_query_history_service.generate_query_embedding(query)
                if embedding:
                    await user_query_history_service.save_query_embedding(
                        db=db, query_history_id=query_history_id,
                        user_id=user_id, embedding=embedding,
                    )
            except Exception as emb_err:
                logger.warning(f"[HISTORY_ENRICH] embedding failed for query {query_history_id}: {emb_err}")
            await db.commit()
            break
    except Exception as e:
        logger.warning(f"[HISTORY_ENRICH] enrichment task failed for query {query_history_id}: {e}")


def set_sql_agent_instance(instance: Optional[object], available: bool = False):
    """Set the SQL agent instance and availability flag."""
    global _sql_agent_instance, _sql_agent_available
    _sql_agent_instance = instance
    _sql_agent_available = available


def get_sql_agent_instance():
    """Get the SQL agent instance."""
    return _sql_agent_instance


def is_sql_agent_available():
    """Check if SQL agent is available."""
    return _sql_agent_available


async def log_chatbot_query(
    user_id: int,
    username: str,
    query: str,
    response: Optional[str],
    success: bool,
    error_message: Optional[str] = None,
    processing_time_ms: Optional[float] = None,
    session_id: Optional[str] = None
):
    """
    Log chatbot query and response to audit log.
    This runs asynchronously to avoid blocking the main request.
    """
    try:
        from db_models import ChatbotAuditLog
        from db_connection import get_db
        
        async for db in get_db():
            audit_log = ChatbotAuditLog(
                user_id=user_id,
                username=username,
                query=query,
                response=response,
                success=success,
                error_message=error_message,
                processing_time_ms=processing_time_ms,
                session_id=session_id
            )
            db.add(audit_log)
            await db.commit()
            logger.debug(f"[AUDIT] Logged chatbot query for user {username} (success: {success})")
            break
    except Exception as e:
        # Don't fail the main request if logging fails
        logger.warning(f"[AUDIT] Failed to log chatbot query: {e}")


async def block_user_for_forbidden_sql(
    user_id: int,
    username: str,
    sql_query: str,
    reason: str
):
    """
    Block a user from using the system when they attempt forbidden SQL operations.
    This runs asynchronously to avoid blocking the main request.
    Also notifies all admins about the security incident.
    """
    try:
        from backend.services.user_service import UserService
        from db_connection import get_db
        
        async for db in get_db():
            blocked_reason = f"Attempted forbidden SQL operation: {reason}. Query: {sql_query[:200]}"
            blocked_user = await UserService.block_user(
                user_id=user_id,
                reason=blocked_reason,
                db=db
            )
            logger.error(f"[SECURITY] ⚠️ BLOCKED USER: {username} (ID: {user_id}) for forbidden SQL: {reason}")
            logger.error(f"[SECURITY] Blocked query: {sql_query[:200]}")
            
            # Notify all admins about the security incident
            try:
                from sqlalchemy import select
                from db_models import User
                admin_result = await db.execute(
                    select(User).where(User.role == "admin")
                )
                admins = admin_result.scalars().all()
                
                for admin in admins:
                    logger.error(f"[SECURITY] 🔔 ADMIN NOTIFICATION: User {username} (ID: {user_id}) was blocked for attempting forbidden SQL operation: {reason}")
                    logger.error(f"[SECURITY] Admin {admin.username} (ID: {admin.id}) should review this incident in the admin panel")
                
                # NOTE: Audit log is NOT created here to avoid duplicates
                # The audit log is created in the main query endpoint after security checks
                # This function only handles blocking and admin notifications
            except Exception as notify_error:
                logger.error(f"[SECURITY] Failed to notify admins: {notify_error}", exc_info=True)
            
            break
    except Exception as e:
        # Log but don't fail - blocking is important but shouldn't break the error response
        logger.error(f"[SECURITY] Failed to block user {username}: {e}", exc_info=True)


@router.post("/query")
async def sql_agent_query(
    request: dict,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None
):
    """
    Query the SQL Intelligence Agent with natural language.
    Uses user-specific persistent conversation memory (like ChatGPT).
    
    Request body:
    {
        "query": "Track Joey"
    }
    
    Returns:
    {
        "response": "SURVEILLANCE INTELLIGENCE REPORT...",
        "session_id": "session_id",
        "success": true
    }
    """
    logger.info(f"[SQL_AGENT_API] Query endpoint called. Request type: {type(request)}, Auth available: {AUTH_AVAILABLE}, Current user: {current_user.username if current_user else None}")
    
    if not _sql_agent_available:
        logger.warning("[SQL_AGENT_API] Query request received but SQL_AGENT_AVAILABLE is False")
        return JSONResponse(
            status_code=503,
            content={
                "success": False,
                "error": "SQL Intelligence Agent module is not available. Please check server logs.",
                "response": None
            }
        )
    
    if _sql_agent_instance is None:
        logger.error("[SQL_AGENT_API] Query request received but sql_agent_instance is None")
        return JSONResponse(
            status_code=503,
            content={
                "success": False,
                "error": "SQL Intelligence Agent is not initialized. Please check server logs.",
                "response": None
            }
        )
    
    # Verify agent is functional
    try:
        if not hasattr(_sql_agent_instance, 'query'):
            logger.error("[SQL_AGENT_API] sql_agent_instance missing 'query' method")
            return JSONResponse(
                status_code=503,
                content={
                    "success": False,
                    "error": "SQL Intelligence Agent is not properly configured.",
                    "response": None
                }
            )
    except Exception as check_error:
        logger.error(f"[SQL_AGENT_API] Error checking agent: {str(check_error)}", exc_info=True)
        return JSONResponse(
            status_code=503,
            content={
                "success": False,
                "error": f"SQL Intelligence Agent health check failed: {str(check_error)}",
                "response": None
            }
        )
    
    try:
        # Validate request format
        if not isinstance(request, dict):
            logger.warning("[SQL_AGENT_API] Invalid request format - not a dict")
            return JSONResponse(
                status_code=400,
                content={
                    "success": False,
                    "error": "Invalid request format. Expected JSON object with 'query' field.",
                    "response": None
                }
            )
        
        query = request.get("query", "").strip()
        if not query:
            logger.warning("[SQL_AGENT_API] Empty query received")
            return JSONResponse(
                status_code=400,
                content={
                    "success": False,
                    "error": "Query is required",
                    "response": None
                }
            )
        
        logger.info(f"[SQL_AGENT_API] Processing query: {query[:100]}...")
        start_time = asyncio.get_event_loop().time()
        
        # Get or create user-specific agent instance for persistent memory
        agent_instance = _sql_agent_instance
        user_id = None
        username = None
        if AUTH_AVAILABLE and current_user:
            user_id = current_user.id
            username = current_user.username
            # Check if user is already blocked — structured code, safe message,
            # no internal reason leaked to the browser
            if not current_user.is_active or not current_user.can_use_chatbot:
                return JSONResponse(
                    status_code=403,
                    content=_error_body(
                        "ACCOUNT_BLOCKED",
                        "Your access to the SQL assistant is temporarily restricted. "
                        "Please contact an administrator.",
                    ),
                )

            agent_instance = _get_or_create_user_agent(user_id)

        # Track if security violation was detected (to prevent duplicate audit logs)
        security_violation_detected = False

        # Run query in thread pool with timeout to avoid blocking
        try:
            QUERY_TIMEOUT = SQL_AGENT_TOTAL_TIMEOUT
            try:
                # Concurrency cap: bounded number of simultaneous agent queries;
                # reject with "busy" instead of piling threads onto Ollama.
                try:
                    await asyncio.wait_for(_sql_agent_semaphore.acquire(), timeout=_SEMAPHORE_WAIT_SECONDS)
                except asyncio.TimeoutError:
                    return JSONResponse(
                        status_code=503,
                        content={"success": False, "error": _BUSY_MESSAGE, "response": None},
                    )

                user_lock = _get_user_lock(user_id) if user_id else None
                lock_acquired = False
                try:
                    if user_lock:
                        await user_lock.acquire()
                        lock_acquired = True
                    logger.info(f"[SQL_AGENT_API] Starting query with {QUERY_TIMEOUT}s timeout: {query[:100]}...")
                    agent_result = await asyncio.wait_for(
                        run_in_threadpool(agent_instance.query, query),
                        timeout=QUERY_TIMEOUT
                    )
                finally:
                    if lock_acquired:
                        user_lock.release()
                    _sql_agent_semaphore.release()
            except asyncio.TimeoutError:
                logger.error(f"[SQL_AGENT_API] Query timeout after {QUERY_TIMEOUT} seconds: {query[:100]}")
                return JSONResponse(
                    status_code=504,
                    content={
                        "success": False,
                        "error": f"Query took too long to process (over {QUERY_TIMEOUT} seconds). Please try a simpler question or check if the system is busy.",
                        "response": None
                    }
                )
            
            # Handle agent response (may be string or tuple with security flags).
            # PRIVACY: never log the response body — lengths and status only.
            if isinstance(agent_result, tuple):
                response, result_dict = agent_result
                logger.info(f"[SQL_AGENT_API] Agent returned tuple response ({len(response)} chars)")
            else:
                response = agent_result
                result_dict = {}
                logger.info(f"[SQL_AGENT_API] Agent returned string response ({len(response)} chars)")
            
            # SECURITY: only an EXPLICIT agent flag counts as a violation.
            # The old "Layer 2" that substring-scanned the model's PROSE for
            # words like "blocked"/"read-only" and then blocked the account is
            # deliberately removed: a denial explanation is not an attack.
            # Policy: Denied -> Audited -> Explained safely; ACCOUNT_BLOCKED
            # only after the explicit threshold in _handle_security_denial.
            if result_dict.get("security_block_user"):
                block_reason = result_dict.get("security_block_reason", "Attempted forbidden SQL operation")
                execution_time_ms = (asyncio.get_event_loop().time() - start_time) * 1000
                denial = await _handle_security_denial(
                    current_user, query, block_reason, execution_time_ms,
                    session_id=agent_instance.conversation_memory.current_session_id
                    if hasattr(agent_instance, "conversation_memory") else None,
                )
                security_violation_detected = True
                return JSONResponse(status_code=403, content=denial)
        except Exception as agent_error:
            logger.error(f"[SQL_AGENT_API] Agent execution error: {str(agent_error)}", exc_info=True)
            error_message = str(agent_error)

            # NOTE: exceptions never auto-block. Make messages user-safe.
            if "timeout" in error_message.lower():
                error_message = "Query timed out. Please try a simpler question."
            elif "connection" in error_message.lower() or "database" in error_message.lower():
                error_message = "Database connection error. Please try again later."
            else:
                error_message = "The assistant could not process this query. Please try again."
            
            return JSONResponse(
                status_code=500,
                content={
                    "success": False,
                    "error": f"Error processing query: {error_message}",
                    "response": None
                }
            )
        
        execution_time = asyncio.get_event_loop().time() - start_time
        execution_time_ms = execution_time * 1000  # Convert to milliseconds
        
        # Validate response
        if not response or not isinstance(response, str):
            logger.error(f"[SQL_AGENT_API] Invalid response from agent: {type(response)}")
            
            # Log failed query to audit log
            if AUTH_AVAILABLE and current_user:
                await log_chatbot_query(
                    user_id=current_user.id,
                    username=current_user.username,
                    query=query,
                    response=None,
                    success=False,
                    error_message="Agent returned invalid response format",
                    processing_time_ms=execution_time_ms,
                    session_id=agent_instance.conversation_memory.current_session_id
                )
            
            return JSONResponse(
                status_code=500,
                content={
                    "success": False,
                    "error": "Agent returned invalid response format",
                    "response": None
                }
            )
        
        # Get current session ID from the agent instance used
        session_id = agent_instance.conversation_memory.current_session_id
        
        # PRIVACY: log status + length only — never the response body
        logger.info(f"[SQL_AGENT_API] ✅ Query completed in {execution_time:.2f}s ({len(response)} chars)")
        
        # Log successful query to audit log (only if no security violation was detected)
        # Security violations are logged in the security check blocks above
        if AUTH_AVAILABLE and current_user and not security_violation_detected:
            await log_chatbot_query(
                user_id=current_user.id,
                username=current_user.username,
                query=query,
                response=response,
                success=True,
                error_message=None,
                processing_time_ms=execution_time_ms,
                session_id=session_id
            )
            
            # Save to user query history so it appears in the sidebar.
            # Awaited (not fire-and-forget) so it reliably commits before this
            # request returns; the core insert is fast (embeddings run in the
            # background) so the added latency is negligible.
            if AUTH_AVAILABLE and current_user:
                metadata = {}
                if result_dict:
                    if "generated_sql" in result_dict:
                        metadata["sql"] = result_dict.get("generated_sql")
                    if "intent" in result_dict:
                        metadata["intent"] = result_dict.get("intent")
                    if "query_result" in result_dict:
                        result_data = result_dict.get("query_result", {})
                        if isinstance(result_data, dict):
                            metadata["row_count"] = result_data.get("row_count", 0)

                await persist_query_history(
                    user_id=current_user.id,
                    query=query,
                    response=response,
                    session_id=session_id,
                    success=True,
                    processing_time_ms=execution_time_ms,
                    metadata=metadata,
                )

        return {
            "success": True,
            "response": response,
            "session_id": session_id,
            "timestamp": datetime.utcnow().isoformat()
        }
    
    except Exception as e:
        logger.error(f"[SQL_AGENT_API] Unexpected error: {str(e)}", exc_info=True)
        return JSONResponse(
            status_code=500,
            content={
                "success": False,
                "error": f"Unexpected error: {str(e)}",
                "response": None
            }
        )


@router.post("/query/stream")
async def sql_agent_query_stream(
    request: dict,
    http_request: Request,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None
):
    """
    Query the SQL Intelligence Agent with streaming SSE response.
    
    Request body:
    {
        "query": "Track Joey",
        "session_id": "optional_session_id"
    }
    
    Returns:
    Server-Sent Events stream with progress updates and final response
    """
    if not _sql_agent_available or _sql_agent_instance is None:
        async def error_stream():
            yield f"data: {json.dumps({'type': 'error', 'message': 'SQL Agent not available'})}\n\n"
        return StreamingResponse(error_stream(), media_type="text/event-stream")
    
    try:
        query = request.get("query", "").strip()
        request_id = _normalize_request_id(request.get("request_id"))

        def _single_event_stream(payload: dict):
            async def _stream():
                yield _sse_event(payload, request_id, 0)
                yield _sse_event({"type": "complete", "success": False}, request_id, 1)
            return StreamingResponse(_stream(), media_type="text/event-stream",
                                     headers={"Cache-Control": "no-cache, no-transform",
                                              "X-Accel-Buffering": "no"})

        if not query:
            return _single_event_stream({"type": "error",
                                         "error_code": "INVALID_REQUEST",
                                         "message": "Query is required"})

        logger.info(f"[SQL_AGENT_API] request_id={request_id} streaming query ({len(query)} chars)")

        # Check if user is already blocked BEFORE processing — structured code,
        # safe message, no internal blocked_reason leaked to the browser
        if AUTH_AVAILABLE and current_user:
            if not current_user.is_active or not current_user.can_use_chatbot:
                return _single_event_stream({
                    "type": "error",
                    "error_code": "ACCOUNT_BLOCKED",
                    "message": "Your access to the SQL assistant is temporarily restricted. "
                               "Please contact an administrator.",
                })

        # Idempotency: the same request_id is executed at most once — a client
        # falling back to another transport cannot run the query twice.
        cancel_event = threading.Event()
        if not _register_request(request_id,
                                 current_user.id if (AUTH_AVAILABLE and current_user) else None,
                                 cancel_event):
            return _single_event_stream({
                "type": "error",
                "error_code": "DUPLICATE_REQUEST",
                "message": "This request was already accepted. Not executing it again.",
            })

        # Get or create user-specific agent instance for persistent memory
        agent_instance = _sql_agent_instance
        if AUTH_AVAILABLE and current_user:
            agent_instance = _get_or_create_user_agent(current_user.id)

        async def stream_query():
            final_response = None
            stream_start_time = asyncio.get_event_loop().time()
            stream_success = False
            completion_sent = False
            last_heartbeat = stream_start_time
            heartbeat_interval = 20.0  # heartbeats keep proxies + client alive
            seq = 0

            def evt(payload: dict) -> str:
                nonlocal seq
                seq += 1
                return _sse_event(payload, request_id, seq)

            # Cancellation + isolation state (cancel_event is registered in the
            # request registry — POST /requests/{id}/cancel sets it server-side)
            sem_acquired = False
            lock_acquired = False
            user_lock = _get_user_lock(current_user.id) if (AUTH_AVAILABLE and current_user) else None
            deadline = stream_start_time + SQL_AGENT_TOTAL_TIMEOUT
            was_cancelled = False

            try:
                # Concurrency cap: bounded simultaneous agent queries
                try:
                    await asyncio.wait_for(_sql_agent_semaphore.acquire(), timeout=_SEMAPHORE_WAIT_SECONDS)
                    sem_acquired = True
                except asyncio.TimeoutError:
                    yield evt({"type": "error", "error_code": "AGENT_BUSY", "message": _BUSY_MESSAGE})
                    yield evt({"type": "complete", "success": False})
                    completion_sent = True
                    return

                # Per-user serialization (double-submit / second tab)
                if user_lock is not None:
                    await user_lock.acquire()
                    lock_acquired = True

                # TRUE streaming: dedicated thread pumps updates into an asyncio
                # queue; each update is forwarded to the client the moment the
                # agent produces it (first byte during generation, not after).
                update_queue = _start_stream_thread(
                    agent_instance, query, cancel_event, asyncio.get_running_loop()
                )

                while True:
                    now = asyncio.get_event_loop().time()
                    if now >= deadline:
                        cancel_event.set()
                        logger.warning(f"[SQL_AGENT_API] request_id={request_id} stream timeout after {SQL_AGENT_TOTAL_TIMEOUT}s")
                        yield evt({"type": "error", "error_code": "QUERY_TIMEOUT",
                                   "message": "Query took too long and was cancelled.", "retryable": True})
                        yield evt({"type": "complete", "success": False})
                        completion_sent = True
                        stream_success = False
                        break

                    # Server-side cancellation (POST /requests/{id}/cancel or WS
                    # cancel): terminal 'cancelled' event, LLM/DB thread stops
                    # at its next boundary via the shared cancel_event.
                    if cancel_event.is_set():
                        was_cancelled = True
                        yield evt({"type": "cancelled"})
                        completion_sent = True
                        stream_success = False
                        break

                    try:
                        update = await asyncio.wait_for(update_queue.get(), timeout=2.0)
                    except asyncio.TimeoutError:
                        # No update yet (LLM thinking): check disconnect + heartbeat
                        try:
                            if await http_request.is_disconnected():
                                cancel_event.set()
                                logger.info(f"[SQL_AGENT_API] request_id={request_id} client disconnected - cancelling")
                                return
                        except Exception:
                            pass
                        current_time = asyncio.get_event_loop().time()
                        if current_time - last_heartbeat >= heartbeat_interval:
                            yield evt({"type": "heartbeat",
                                       "timestamp": datetime.utcnow().isoformat() + "Z"})
                            last_heartbeat = current_time
                        continue

                    if update is _STREAM_SENTINEL:
                        break

                    try:
                        # Format as SSE (request_id + sequence on every event)
                        yield evt(update)

                        # Capture final response
                        if update.get("type") == "complete":
                            # Check if response is in completion message (some agents include it)
                            if "response" in update:
                                final_response = update.get("response")
                            # If we don't have final_response yet but have accumulated content, use that
                            elif final_response is None or len(final_response) == 0:
                                # Check if we have accumulated content from previous chunks
                                # This handles cases where content was streamed but not captured
                                logger.debug(f"[SQL_AGENT_API] Completion received but no response field, checking accumulated content")
                            # Check if the completion message indicates success
                            if update.get("success") is not False:  # True or undefined means success
                                stream_success = True
                            else:
                                stream_success = False
                            completion_sent = True
                            logger.debug(f"[SQL_AGENT_API] Completion received: success={stream_success}, final_response_length={len(final_response) if final_response else 0}, response_length={update.get('response_length', 'N/A')}")
                        elif update.get("type") == "content":
                            # Accumulate content chunks - this is how we build the final response
                            if final_response is None:
                                final_response = ""
                            content_chunk = update.get("content", "")
                            final_response += content_chunk
                            logger.debug(f"[SQL_AGENT_API] Content chunk received: {len(content_chunk)} chars, total accumulated: {len(final_response)} chars")
                        elif update.get("type") == "error":
                            stream_success = False
                            final_response = update.get("message", "An error occurred")
                            # A denied operation is Denied -> Audited -> Explained.
                            # "Security:" here is OUR validator's deterministic
                            # prefix (sql_agent/database.py), not model prose.
                            # It records a violation; blocking only happens at
                            # the explicit threshold inside the policy handler.
                            error_message = update.get("message", "")
                            if error_message.startswith("Security:") and AUTH_AVAILABLE and current_user:
                                await _handle_security_denial(
                                    current_user, query, error_message,
                                    (asyncio.get_event_loop().time() - stream_start_time) * 1000,
                                )
                            # Send completion after error to properly close stream
                            yield evt({"type": "complete", "success": False})
                            completion_sent = True
                            break
                    except Exception as yield_error:
                        logger.error(f"[SQL_AGENT_API] Error yielding update: {str(yield_error)}", exc_info=True)
                        try:
                            yield evt({"type": "error", "error_code": "STREAM_ERROR",
                                       "message": "Streaming failed. Please retry.", "retryable": True})
                            yield evt({"type": "complete", "success": False})
                            completion_sent = True
                        except Exception:
                            pass
                        break

                # An explicit terminal event is REQUIRED by the contract —
                # clients never infer completion from content length.
                if not completion_sent:
                    try:
                        yield evt({"type": "complete", "success": stream_success,
                                   "response": final_response[:500] if final_response else None})
                        completion_sent = True
                    except Exception as completion_error:
                        logger.error(f"[SQL_AGENT_API] Error sending completion: {str(completion_error)}")

            except Exception as e:
                logger.error(f"[SQL_AGENT_API] Streaming error (outer): {str(e)}", exc_info=True)
                if not completion_sent:
                    try:
                        yield evt({"type": "error", "error_code": "STREAM_ERROR",
                                   "message": "Streaming failed. Please retry.", "retryable": True})
                        yield evt({"type": "complete", "success": False})
                        completion_sent = True
                    except Exception as close_error:
                        logger.error(f"[SQL_AGENT_API] Error closing stream: {str(close_error)}")
                stream_success = False
                final_response = None
            finally:
                # Stop the agent thread at its next node boundary and free the slot
                # (also runs on client disconnect / GeneratorExit).
                cancel_event.set()
                _finish_request(request_id,
                                "cancelled" if was_cancelled
                                else ("completed" if stream_success else "failed"))
                if lock_acquired and user_lock is not None:
                    try:
                        user_lock.release()
                    except RuntimeError:
                        pass
                if sem_acquired:
                    _sql_agent_semaphore.release()

                if not completion_sent:
                    try:
                        yield evt({"type": "complete", "success": stream_success})
                        completion_sent = True
                    except Exception:
                        pass

                # PRIVACY: structured status only — never the response body
                logger.info(
                    "[SQL_AGENT_API] request_id=%s stream finished status=%s duration=%.1fs response_chars=%d",
                    request_id,
                    "cancelled" if was_cancelled else ("completed" if stream_success else "failed"),
                    asyncio.get_event_loop().time() - stream_start_time,
                    len(final_response) if final_response else 0,
                )

                # Conversation-memory fallback for history persistence only
                if not final_response:
                    try:
                        if hasattr(agent_instance, 'conversation_memory') and agent_instance.conversation_memory:
                            recent_messages = agent_instance.conversation_memory.get_recent_messages(limit=5)
                            for msg in reversed(recent_messages or []):
                                if msg.get('role') == 'assistant' and msg.get('content'):
                                    final_response = msg.get('content')
                                    break
                    except Exception:
                        pass
            
            # Log to audit and save history after streaming completes.
            # Note: do NOT require final_response here - even queries that produced an
            # empty/failed response must still appear in the user's sidebar history.
            if AUTH_AVAILABLE and current_user:
                try:
                    stream_time_ms = (asyncio.get_event_loop().time() - stream_start_time) * 1000
                    session_id = agent_instance.conversation_memory.current_session_id
                    
                    logger.info(f"[SQL_AGENT_API] 📊 Logging audit: user={current_user.username}, success={stream_success}, time={stream_time_ms:.2f}ms, response_length={len(final_response) if final_response else 0}")
                    await log_chatbot_query(
                        user_id=current_user.id,
                        username=current_user.username,
                        query=query,
                        response=final_response,
                        success=stream_success,
                        error_message=None if stream_success else "Streaming error",
                        processing_time_ms=stream_time_ms,
                        session_id=session_id
                    )
                    
                    # Save to user query history so it appears in the sidebar.
                    # Awaited (not fire-and-forget) so it reliably commits within the
                    # request lifecycle. The full response is already streamed to the
                    # client at this point, so awaiting adds no user-facing latency.
                    await persist_query_history(
                        user_id=current_user.id,
                        query=query,
                        response=final_response,
                        session_id=session_id,
                        success=stream_success,
                        processing_time_ms=stream_time_ms,
                        metadata={},
                    )

                except Exception as audit_error:
                    logger.error(f"[SQL_AGENT_API] Error logging audit: {str(audit_error)}", exc_info=True)
        
        return StreamingResponse(
            stream_query(), 
            media_type="text/event-stream; charset=utf-8",
            headers={
                "Cache-Control": "no-cache, no-transform",
            "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
                "Transfer-Encoding": "chunked"
            }
        )
    
    except Exception as e:
        logger.error(f"[SQL_AGENT_API] Stream endpoint error: {str(e)}", exc_info=True)
        async def error_stream():
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
        return StreamingResponse(error_stream(), media_type="text/event-stream")


async def sql_agent_websocket(websocket: WebSocket):
    """
    WebSocket endpoint for SQL Agent queries with real-time streaming.
    This is registered separately in backend.main at /ws/sql-agent
    Uses user-specific persistent conversation memory.
    """
    # Origin validation BEFORE accept: cross-site pages cannot forge Origin
    origin = websocket.headers.get("origin")
    if origin:
        try:
            from urllib.parse import urlparse
            origin_host = urlparse(origin).hostname
            host_header = (websocket.headers.get("host") or "").split(":")[0]
            if origin_host not in {host_header, "localhost", "127.0.0.1"}:
                logger.warning(f"[SQL_AGENT_WS] Rejected cross-origin connection from {origin_host}")
                await websocket.close(code=1008, reason="Origin not allowed")
                return
        except Exception:
            await websocket.close(code=1008, reason="Origin not allowed")
            return

    await websocket.accept()
    logger.info("[SQL_AGENT_WS] WebSocket connection established")

    # Check authentication and get current user. The browser sends the
    # HttpOnly session cookie on the handshake — parse it FIRST (the old
    # token-only path meant cookie-authenticated browsers could never
    # use this transport at all).
    current_user = None
    if AUTH_AVAILABLE:
        try:
            from backend.auth.auth_service import AuthService
            token = None
            cookie_header = websocket.headers.get("cookie", "")
            if cookie_header:
                for cookie in cookie_header.split(";"):
                    cookie = cookie.strip()
                    if cookie.startswith("access_token="):
                        token = cookie.split("=", 1)[1].strip()
                        break
            if not token:
                token = websocket.query_params.get("token") or websocket.headers.get("Authorization", "").replace("Bearer ", "")
            if token:
                payload = AuthService.decode_token(token)
                if payload:
                    # JWT 'sub' claim is a string, convert to int for user lookup
                    user_id_str = payload.get("sub")
                    if user_id_str:
                        try:
                            user_id = int(user_id_str)
                            from db_connection import get_db
                            async for db in get_db():
                                current_user = await AuthService.get_user_by_id(user_id, db)
                                break  # Exit the async for loop after getting the user
                        except (ValueError, TypeError) as e:
                            logger.error(f"[SQL_AGENT_WS] Invalid user ID in token: {user_id_str}, error: {e}")
                            await websocket.close(code=1008, reason="Invalid token")
                            return
                        
                        # Check if user is valid after getting from database
                        if not current_user or not current_user.is_active or not current_user.can_use_chatbot:
                            await websocket.close(code=1008, reason="Unauthorized or chatbot access denied")
                            return
                    else:
                        await websocket.close(code=1008, reason="Invalid token")
                        return
                else:
                    await websocket.close(code=1008, reason="Invalid token")
                    return
            else:
                await websocket.close(code=1008, reason="Authentication required")
                return
        except Exception as e:
            logger.error(f"WebSocket auth error: {e}")
            await websocket.close(code=1008, reason="Authentication failed")
            return
    
    if not _sql_agent_available or _sql_agent_instance is None:
        await websocket.send_json({
            "type": "error",
            "message": "SQL Agent not available"
        })
        await websocket.close()
        return
    
    # Get or create user-specific agent instance for persistent memory
    agent_instance = _sql_agent_instance
    if AUTH_AVAILABLE and current_user:
        agent_instance = _get_or_create_user_agent(current_user.id)
    
    try:
        while True:
            # Receive message from client — typed protocol:
            #   {"type":"query","request_id":..., "query":...}
            #   {"type":"cancel","request_id":...}
            # (legacy bare {"query": ...} still accepted)
            data = await websocket.receive_json()
            if not isinstance(data, dict):
                continue

            msg_type = data.get("type") or ("query" if data.get("query") else None)

            if msg_type == "cancel":
                cancel_id = data.get("request_id")
                entry = _ACTIVE_REQUESTS.get(cancel_id) if isinstance(cancel_id, str) else None
                if entry is not None and (not (AUTH_AVAILABLE and current_user)
                                          or entry.get("user_id") in (None, current_user.id)
                                          or current_user.role == "admin"):
                    entry["cancel_event"].set()
                    entry["status"] = "cancelling"
                    logger.info(f"[SQL_AGENT_WS] request_id={cancel_id} cancel via WebSocket")
                continue

            if msg_type != "query":
                continue

            query = (data.get("query") or "").strip()
            request_id = _normalize_request_id(data.get("request_id"))
            seq = 0

            def ws_evt(payload: dict) -> dict:
                nonlocal seq
                seq += 1
                return {**payload, "request_id": request_id, "sequence": seq}

            if not query:
                await websocket.send_json(ws_evt({"type": "error",
                                                  "error_code": "INVALID_REQUEST",
                                                  "message": "Query is required"}))
                continue

            if query.lower() in ["close", "exit", "quit"]:
                await websocket.send_json(ws_evt({"type": "close", "message": "Closing connection"}))
                break

            logger.info(f"[SQL_AGENT_WS] request_id={request_id} processing query ({len(query)} chars)")
            ws_start_time = asyncio.get_event_loop().time()

            cancel_event = threading.Event()
            # Idempotency: never execute the same request_id twice
            if not _register_request(request_id,
                                     current_user.id if (AUTH_AVAILABLE and current_user) else None,
                                     cancel_event):
                await websocket.send_json(ws_evt({
                    "type": "error", "error_code": "DUPLICATE_REQUEST",
                    "message": "This request was already accepted."}))
                continue

            # Send initial status
            await websocket.send_json(ws_evt({
                "type": "status", "message": "Processing query...", "step": "start"}))

            sem_acquired = False
            lock_acquired = False
            was_cancelled = False
            query_success = False
            accumulated_response = ""
            user_lock = _get_user_lock(current_user.id) if (AUTH_AVAILABLE and current_user) else None
            try:
                # Concurrency cap shared with the HTTP endpoints
                try:
                    await asyncio.wait_for(_sql_agent_semaphore.acquire(), timeout=_SEMAPHORE_WAIT_SECONDS)
                    sem_acquired = True
                except asyncio.TimeoutError:
                    await websocket.send_json(ws_evt({"type": "error", "error_code": "AGENT_BUSY",
                                                      "message": _BUSY_MESSAGE}))
                    await websocket.send_json(ws_evt({"type": "complete", "success": False}))
                    _finish_request(request_id, "failed")
                    continue

                if user_lock is not None:
                    await user_lock.acquire()
                    lock_acquired = True

                # TRUE streaming bridge: agent runs in its own thread, updates are
                # forwarded as they arrive — the blocking generator is never
                # iterated on the event loop.
                update_queue = _start_stream_thread(
                    agent_instance, query, cancel_event, asyncio.get_running_loop()
                )
                deadline = asyncio.get_event_loop().time() + SQL_AGENT_TOTAL_TIMEOUT

                accumulated_response = ""
                query_success = False
                last_heartbeat = asyncio.get_event_loop().time()
                while True:
                    if asyncio.get_event_loop().time() >= deadline:
                        cancel_event.set()
                        await websocket.send_json(ws_evt({"type": "error", "error_code": "QUERY_TIMEOUT",
                                                          "message": "Query took too long and was cancelled.",
                                                          "retryable": True}))
                        await websocket.send_json(ws_evt({"type": "complete", "success": False}))
                        break

                    # Server-side cancellation (cancel message or REST endpoint)
                    if cancel_event.is_set():
                        was_cancelled = True
                        await websocket.send_json(ws_evt({"type": "cancelled"}))
                        break

                    try:
                        update = await asyncio.wait_for(update_queue.get(), timeout=2.0)
                    except asyncio.TimeoutError:
                        now = asyncio.get_event_loop().time()
                        if now - last_heartbeat >= 20.0:
                            await websocket.send_json(ws_evt({
                                "type": "heartbeat",
                                "timestamp": datetime.utcnow().isoformat() + "Z"}))
                            last_heartbeat = now
                        continue  # still generating; loop re-checks deadline/cancel

                    if update is _STREAM_SENTINEL:
                        break

                    await websocket.send_json(ws_evt(update))

                    # Capture final response
                    if update.get("type") == "complete":
                        if update.get("response"):
                            accumulated_response = update.get("response")
                        query_success = update.get("success") is not False
                    elif update.get("type") == "content":
                        accumulated_response += update.get("content", "")
                    elif update.get("type") == "error":
                        query_success = False
                        # Denied -> audited -> explained (threshold policy, never
                        # instant blocking; "Security:" is our validator's prefix)
                        error_message = update.get("message", "")
                        if error_message.startswith("Security:") and AUTH_AVAILABLE and current_user:
                            await _handle_security_denial(
                                current_user, query, error_message,
                                (asyncio.get_event_loop().time() - ws_start_time) * 1000)

                    # Break if complete or error
                    if update.get("type") in ["complete", "error"]:
                        break

                # Persist to sidebar history (awaited so it reliably commits).
                # The WebSocket path previously never saved history at all - this was
                # a second reason some queries were missing from the sidebar.
                if AUTH_AVAILABLE and current_user:
                    ws_time_ms = (asyncio.get_event_loop().time() - ws_start_time) * 1000
                    ws_session_id = None
                    try:
                        ws_session_id = agent_instance.conversation_memory.current_session_id
                    except Exception:
                        pass
                    await persist_query_history(
                        user_id=current_user.id,
                        query=query,
                        response=accumulated_response or None,
                        session_id=ws_session_id,
                        success=query_success,
                        processing_time_ms=ws_time_ms,
                        metadata={},
                    )

            except WebSocketDisconnect:
                # Client vanished mid-query: stop the agent thread, then bail out
                cancel_event.set()
                raise
            except Exception as e:
                logger.error(f"[SQL_AGENT_WS] Query processing error: {str(e)}", exc_info=True)
                await websocket.send_json(ws_evt({
                    "type": "error", "error_code": "STREAM_ERROR",
                    "message": "The assistant could not process this query.", "retryable": True}))
            finally:
                cancel_event.set()
                _finish_request(request_id,
                                "cancelled" if was_cancelled
                                else ("completed" if query_success else "failed"))
                if lock_acquired and user_lock is not None:
                    try:
                        user_lock.release()
                    except RuntimeError:
                        pass
                if sem_acquired:
                    _sql_agent_semaphore.release()

    except WebSocketDisconnect:
        logger.info("[SQL_AGENT_WS] WebSocket disconnected")
    except Exception as e:
        logger.error(f"[SQL_AGENT_WS] WebSocket error: {str(e)}", exc_info=True)
        try:
            await websocket.send_json({
                "type": "error",
                "message": f"WebSocket error: {str(e)}"
            })
        except:
            pass
    finally:
        try:
            await websocket.close()
        except:
            pass


@router.post("/requests/{request_id}/cancel")
async def cancel_sql_agent_request(
    request_id: str,
    http_request: Request,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None,
    _csrf: None = Depends(require_sql_agent_csrf),
):
    """Server-side cancellation: stops the LLM/SQL worker thread at its next
    boundary via the request's registered cancel event (the stream then emits
    a terminal 'cancelled' event). Owner or admin only."""
    entry = _ACTIVE_REQUESTS.get(request_id)
    if entry is None:
        return JSONResponse(status_code=404, content=_error_body(
            "REQUEST_NOT_FOUND", "No such request.", retryable=False))

    if AUTH_AVAILABLE and current_user:
        if entry.get("user_id") not in (None, current_user.id) and current_user.role != "admin":
            return JSONResponse(status_code=403, content=_error_body(
                "ACCESS_DENIED", "You cannot cancel another user's request."))

    if entry["status"] != "running":
        return {"success": True, "request_id": request_id, "status": entry["status"]}

    entry["cancel_event"].set()
    entry["status"] = "cancelling"
    logger.info("[SQL_AGENT_API] request_id=%s cancel requested by user_id=%s",
                request_id, current_user.id if (AUTH_AVAILABLE and current_user) else None)
    return {"success": True, "request_id": request_id, "status": "cancelling"}


# Health-result cache: the endpoint must stay LIGHTWEIGHT (no LLM calls, no
# model init, no expensive SQL) — one SELECT 1, cached briefly.
_health_cache = {"at": 0.0, "body": None}
_HEALTH_CACHE_SECONDS = 15.0


@router.get("/health")
async def sql_agent_health(response: Response):
    """Lightweight component health: never invokes the LLM or heavy SQL."""
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate"
    now = asyncio.get_event_loop().time()
    if _health_cache["body"] is not None and now - _health_cache["at"] < _HEALTH_CACHE_SECONDS:
        return _health_cache["body"]

    components = {
        "model": "ready" if (_sql_agent_available and _sql_agent_instance is not None) else "unavailable",
        "database": "unknown",
        "history": "ready" if AUTH_AVAILABLE else "unavailable",
    }

    if _sql_agent_instance is not None:
        try:
            # SELECT 1 on a threadpool with a hard 3s cap — the only I/O here
            db_ok = await asyncio.wait_for(
                run_in_threadpool(_sql_agent_instance.test_connection), timeout=3.0)
            components["database"] = "ready" if db_ok else "error"
        except Exception:
            components["database"] = "error"

    status = "operational" if all(v == "ready" for k, v in components.items()
                                  if k in ("model", "database")) else "degraded"
    body = {
        "available": _sql_agent_available and _sql_agent_instance is not None,
        "status": status,
        "components": components,
        "checked_at": datetime.utcnow().isoformat() + "Z",
    }
    _health_cache["at"] = now
    _health_cache["body"] = body
    return body


@router.get("/schema")
async def sql_agent_schema():
    """Get database schema description"""
    logger.debug("[SQL_AGENT_API] Schema request received")
    
    if not _sql_agent_available or _sql_agent_instance is None:
        logger.warning("[SQL_AGENT_API] Schema request - agent not available")
        raise HTTPException(status_code=503, detail="SQL Agent not available")
    
    try:
        schema = _sql_agent_instance.get_schema()
        logger.info(f"[SQL_AGENT_API] Schema retrieved successfully ({len(schema)} chars)")
        return {
            "success": True,
            "schema": schema
        }
    except Exception as e:
        logger.error(f"[SQL_AGENT_API] Schema retrieval error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/session/new")
async def sql_agent_new_session():
    """Start a new conversation session"""
    logger.info("[SQL_AGENT_API] New session request received")
    
    if not _sql_agent_available or _sql_agent_instance is None:
        logger.warning("[SQL_AGENT_API] New session - agent not available")
        raise HTTPException(status_code=503, detail="SQL Agent not available")
    
    try:
        session_id = _sql_agent_instance.conversation_memory.start_session()
        logger.info(f"[SQL_AGENT_API] New session created: {session_id}")
        return {
            "success": True,
            "session_id": session_id,
            "message": "New session created"
        }
    except Exception as e:
        logger.error(f"[SQL_AGENT_API] New session error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/sessions")
async def sql_agent_list_sessions():
    """List all conversation sessions"""
    logger.debug("[SQL_AGENT_API] List sessions request received")
    
    if not _sql_agent_available or _sql_agent_instance is None:
        logger.warning("[SQL_AGENT_API] List sessions - agent not available")
        raise HTTPException(status_code=503, detail="SQL Agent not available")
    
    try:
        sessions = _sql_agent_instance.conversation_memory.list_sessions()
        logger.info(f"[SQL_AGENT_API] Listed {len(sessions)} sessions")
        return {
            "success": True,
            "sessions": sessions,
            "current_session": _sql_agent_instance.conversation_memory.current_session_id
        }
    except Exception as e:
        logger.error(f"[SQL_AGENT_API] List sessions error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/session/load")
async def sql_agent_load_session(request: dict):
    """Load a previous conversation session"""
    logger.info("[SQL_AGENT_API] Load session request received")
    
    if not _sql_agent_available or _sql_agent_instance is None:
        logger.warning("[SQL_AGENT_API] Load session - agent not available")
        raise HTTPException(status_code=503, detail="SQL Agent not available")
    
    try:
        session_id = request.get("session_id")
        if not session_id:
            logger.warning("[SQL_AGENT_API] Load session - session_id missing")
            raise HTTPException(status_code=400, detail="session_id is required")
        
        logger.debug(f"[SQL_AGENT_API] Attempting to load session: {session_id}")
        loaded = _sql_agent_instance.conversation_memory.load_session(session_id)
        if loaded:
            logger.info(f"[SQL_AGENT_API] Session loaded successfully: {session_id}")
            return {
                "success": True,
                "session_id": session_id,
                "message": "Session loaded successfully"
            }
        else:
            logger.warning(f"[SQL_AGENT_API] Session not found: {session_id}")
            raise HTTPException(status_code=404, detail="Session not found")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[SQL_AGENT_API] Load session error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# Export Models — server-enforced size limits (Pydantic rejects oversize)
class ExportRequest(BaseModel):
    content: str  # validated in _sanitize_export below
    title: str
    timestamp: str


_EXPORT_MAX_CONTENT_CHARS = 500_000
_EXPORT_MAX_TITLE_CHARS = 200
# Bounded concurrent document generation (reportlab/docx are CPU-bound)
_export_semaphore = asyncio.Semaphore(4)


def _sanitize_export(request: "ExportRequest"):
    """Size limits + markup-injection prevention for document generation.

    reportlab's Paragraph parses XML-ish markup — raw '<'/'&' from the model
    or browser must be escaped BEFORE we selectively re-allow <b>/<i>.
    Returns (safe_title, safe_content, safe_date) or raises 413/422.
    """
    from xml.sax.saxutils import escape as _xml_escape
    if len(request.content) > _EXPORT_MAX_CONTENT_CHARS:
        raise HTTPException(status_code=413, detail="Export content too large")
    title = re.sub(r'[<>&\x00-\x1f]', '', request.title or 'Intelligence Report').strip()
    title = title[:_EXPORT_MAX_TITLE_CHARS] or 'Intelligence Report'
    content = _xml_escape(request.content)
    # Safe date for the filename (never trust raw client timestamp strings)
    safe_date = datetime.utcnow().strftime('%Y-%m-%d')
    return title, content, safe_date


@router.post("/export/pdf")
async def export_to_pdf(
    request: ExportRequest,
    http_request: Request,
    current_user: User = Depends(require_chatbot_access()),
    _csrf: None = Depends(require_sql_agent_csrf),
):
    """Export chatbot response to PDF format (auth + CSRF + size limits +
    markup-injection prevention + bounded concurrency + audit log line)."""
    safe_title, safe_content, safe_date = _sanitize_export(request)
    logger.info("[EXPORT] format=pdf user_id=%s content_chars=%d",
                current_user.id if current_user else None, len(request.content))
    async with _export_semaphore:
        return await asyncio.wait_for(
            run_in_threadpool(_build_pdf_export, safe_title, safe_content, safe_date,
                              current_user.username if current_user else 'System'),
            timeout=60.0,
        )


def _build_pdf_export(safe_title: str, safe_content: str, safe_date: str, analyst: str):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        # Create PDF in memory
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, 
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=72)
        
        # Container for PDF content
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#00ff96'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#00ff96'),
            spaceAfter=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.black,
            leading=14,
            spaceAfter=12
        )
        
        # Add title
        title = Paragraph(f"<b>INTELLIGENCE REPORT</b>", title_style)
        story.append(title)
        story.append(Spacer(1, 0.2*inch))

        # Add metadata (all values pre-sanitized in _sanitize_export)
        metadata = f"<b>Query:</b> {safe_title}<br/>"
        metadata += f"<b>Generated:</b> {safe_date}<br/>"
        metadata += f"<b>Analyst:</b> {re.sub(r'[<>&]', '', analyst)}"
        story.append(Paragraph(metadata, body_style))
        story.append(Spacer(1, 0.3*inch))

        # Content arrives XML-escaped; re-allow ONLY escaped bold/italic markers
        content = safe_content
        content = content.replace('&lt;br&gt;', '\n').replace('&lt;br/&gt;', '\n')
        content = re.sub(r'&lt;strong&gt;(.*?)&lt;/strong&gt;', r'<b>\1</b>', content)
        content = re.sub(r'&lt;em&gt;(.*?)&lt;/em&gt;', r'<i>\1</i>', content)

        # Split into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Clean up the paragraph
                para = para.strip().replace('\n', '<br/>')
                story.append(Paragraph(para, body_style))
                story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF bytes
        buffer.seek(0)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        # Return PDF as response — filename is server-built, never client input
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="Intelligence_Report_{safe_date}.pdf"'
            }
        )
    except ImportError:
        logger.error("[EXPORT] reportlab not installed. Install with: pip install reportlab")
        raise HTTPException(
            status_code=500,
            detail="PDF export is not available on this server."
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[EXPORT] PDF export error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate PDF")


@router.post("/export/word")
async def export_to_word(
    request: ExportRequest,
    http_request: Request,
    current_user: User = Depends(require_chatbot_access()),
    _csrf: None = Depends(require_sql_agent_csrf),
):
    """Export chatbot response to Word (auth + CSRF + size limits + bounded
    concurrency; docx runs are plain text — no markup injection possible)."""
    safe_title, _safe_content, safe_date = _sanitize_export(request)
    logger.info("[EXPORT] format=word user_id=%s content_chars=%d",
                current_user.id if current_user else None, len(request.content))
    async with _export_semaphore:
        return await asyncio.wait_for(
            run_in_threadpool(_build_word_export, safe_title, request.content, safe_date,
                              current_user.username if current_user else 'System'),
            timeout=60.0,
        )


def _build_word_export(safe_title: str, raw_content: str, safe_date: str, analyst: str):
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create Word document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('INTELLIGENCE REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(0, 255, 150)  # #00ff96
        
        # Add metadata (sanitized title, server date; docx runs are plain text)
        doc.add_paragraph(f'Query: {safe_title}')
        doc.add_paragraph(f'Generated: {safe_date}')
        doc.add_paragraph(f'Analyst: {analyst}')
        doc.add_paragraph()  # Empty line

        # Add content
        content = raw_content
        # Clean HTML
        content = re.sub(r'<br\s*/?>', '\n', content)
        content = re.sub(r'<strong>(.*?)</strong>', r'\1', content)  # Remove strong tags
        content = re.sub(r'<em>(.*?)</em>', r'\1', content)  # Remove em tags
        content = re.sub(r'<[^>]+>', '', content)  # Remove remaining HTML tags
        
        # Split into paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.style.font.size = Pt(11)
                p.style.font.name = 'Calibri'
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        word_bytes = buffer.getvalue()
        buffer.close()
        
        # Return Word document — filename is server-built, never client input
        return Response(
            content=word_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="Intelligence_Report_{safe_date}.docx"'
            }
        )
    except ImportError:
        logger.error("[EXPORT] python-docx not installed. Install with: pip install python-docx")
        raise HTTPException(
            status_code=500,
            detail="Word export is not available on this server."
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[EXPORT] Word export error: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to generate Word document")


# =====================================================
# QUERY HISTORY ENDPOINTS
# =====================================================

@router.get("/history")
async def get_query_history(
    response: Response,
    page: int = 1,
    page_size: int = 25,
    limit: Optional[int] = None,   # legacy compat: maps onto page_size
    offset: Optional[int] = None,  # legacy compat: maps onto page
    session_id: Optional[str] = None,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None
):
    """User's query history — server-side pagination, owner-scoped.

    Returns {history, page, page_size, count, has_more}. Timestamps are
    timezone-aware ISO-8601 (UTC, 'Z' suffix).
    """
    if not AUTH_AVAILABLE or not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")

    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate"

    if limit is not None:
        page_size = limit
    page_size = max(1, min(int(page_size), 100))
    page = max(1, int(page))
    computed_offset = offset if offset is not None else (page - 1) * page_size

    try:
        async for db in get_db():
            # Fetch one extra row to compute has_more without a COUNT(*)
            history = await user_query_history_service.get_user_query_history(
                db=db,
                user_id=current_user.id,
                limit=page_size + 1,
                offset=computed_offset,
                session_id=session_id
            )
            has_more = len(history) > page_size
            history = history[:page_size]

            return {
                "success": True,
                "history": [
                    {
                        "id": q.id,
                        "query": q.query_text,
                        "response": q.response_text[:500] if q.response_text else None,
                        "timestamp": q.query_timestamp.isoformat() + "Z",
                        "success": q.success,
                        "processing_time_ms": q.processing_time_ms,
                        "session_id": q.session_id,
                        "metadata": q.query_metadata or {}
                    }
                    for q in history
                ],
                "page": page,
                "page_size": page_size,
                "count": len(history),
                "has_more": has_more,
            }
    except Exception as e:
        logger.error(f"[HISTORY] Error getting query history: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to load history")


@router.get("/history/{query_id}")
async def get_query_by_id(
    query_id: int,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None
):
    """Get a specific query by ID."""
    if not AUTH_AVAILABLE or not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    try:
        async for db in get_db():
            query = await user_query_history_service.get_query_by_id(
                db=db,
                query_id=query_id,
                user_id=current_user.id
            )
            
            if not query:
                raise HTTPException(status_code=404, detail="Query not found")
            
            return {
                "success": True,
                "query": {
                    "id": query.id,
                    "query": query.query_text,
                    "response": query.response_text,
                    "timestamp": query.query_timestamp.isoformat(),
                    "response_timestamp": query.response_timestamp.isoformat() if query.response_timestamp else None,
                    "success": query.success,
                    "error_message": query.error_message,
                    "processing_time_ms": query.processing_time_ms,
                    "session_id": query.session_id,
                    "metadata": query.query_metadata or {}  # Fixed: use query_metadata instead of metadata_
                }
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[HISTORY] Error getting query: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/history/{query_id}")
async def delete_query_from_history(
    query_id: int,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None,
    _csrf: None = Depends(require_sql_agent_csrf),
):
    """Permanently delete one of the CALLER'S OWN history entries (the service
    filters by user_id — object-level authorization, no ID guessing). The
    chatbot audit log entry is preserved as the security record."""
    if not AUTH_AVAILABLE or not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")

    try:
        async for db in get_db():
            deleted = await user_query_history_service.delete_query(
                db=db,
                query_id=query_id,
                user_id=current_user.id
            )

            if not deleted:
                raise HTTPException(status_code=404, detail="Query not found")

            return {"success": True, "deleted_id": query_id}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[HISTORY] Error deleting query: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# =====================================================
# SESSION MANAGEMENT ENDPOINTS
# =====================================================

@router.get("/sessions/list")
async def list_user_sessions(
    active_only: bool = False,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None
):
    """List user's conversation sessions."""
    if not AUTH_AVAILABLE or not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    try:
        async for db in get_db():
            sessions = await user_query_history_service.get_user_sessions(
                db=db,
                user_id=current_user.id,
                active_only=active_only
            )
            
            return {
                "success": True,
                "sessions": [
                    {
                        "session_id": s.session_id,
                        "session_name": s.session_name,
                        "started_at": s.started_at.isoformat(),
                        "last_activity_at": s.last_activity_at.isoformat(),
                        "is_active": s.is_active,
                        "query_count": s.query_count,
                        "context_summary": s.context_summary
                    }
                    for s in sessions
                ]
            }
    except Exception as e:
        logger.error(f"[SESSIONS] Error listing sessions: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/sessions/create")
async def create_session(
    request: dict,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None
):
    """Create a new conversation session."""
    if not AUTH_AVAILABLE or not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    try:
        session_name = request.get("session_name")
        session_id = request.get("session_id")
        
        async for db in get_db():
            session = await user_query_history_service.get_or_create_session(
                db=db,
                user_id=current_user.id,
                session_id=session_id,
                session_name=session_name
            )
            await db.commit()
            
            return {
                "success": True,
                "session": {
                    "session_id": session.session_id,
                    "session_name": session.session_name,
                    "started_at": session.started_at.isoformat(),
                    "last_activity_at": session.last_activity_at.isoformat()
                }
            }
    except Exception as e:
        logger.error(f"[SESSIONS] Error creating session: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/sessions/{session_id}")
async def update_session(
    session_id: str,
    request: dict,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None
):
    """Update a session (e.g., context summary, active status)."""
    if not AUTH_AVAILABLE or not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    try:
        context_summary = request.get("context_summary")
        is_active = request.get("is_active")
        
        async for db in get_db():
            session = await user_query_history_service.update_session(
                db=db,
                user_id=current_user.id,
                session_id=session_id,
                context_summary=context_summary,
                is_active=is_active
            )
            await db.commit()
            
            if not session:
                raise HTTPException(status_code=404, detail="Session not found")
            
            return {
                "success": True,
                "session": {
                    "session_id": session.session_id,
                    "session_name": session.session_name,
                    "context_summary": session.context_summary,
                    "is_active": session.is_active,
                    "last_activity_at": session.last_activity_at.isoformat()
                }
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[SESSIONS] Error updating session: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# =====================================================
# MEMORY MANAGEMENT ENDPOINTS
# =====================================================

@router.get("/memory")
async def get_user_memories(
    memory_type: Optional[str] = None,
    min_importance: int = 0,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None
):
    """Get user's memories."""
    if not AUTH_AVAILABLE or not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    try:
        from db_models import MemoryType
        
        mem_type = None
        if memory_type:
            try:
                mem_type = MemoryType(memory_type)
            except ValueError:
                raise HTTPException(status_code=400, detail=f"Invalid memory type: {memory_type}")
        
        async for db in get_db():
            memories = await user_query_history_service.get_user_memories(
                db=db,
                user_id=current_user.id,
                memory_type=mem_type,
                min_importance=min_importance
            )
            
            return {
                "success": True,
                "memories": [
                    {
                        "id": m.id,
                        "type": m.memory_type.value,
                        "key": m.memory_key,
                        "value": m.memory_value,
                        "importance": m.importance_score,
                        "created_at": m.created_at.isoformat(),
                        "last_accessed_at": m.last_accessed_at.isoformat() if m.last_accessed_at else None,
                        "access_count": m.access_count,
                        "expires_at": m.expires_at.isoformat() if m.expires_at else None
                    }
                    for m in memories
                ],
                "count": len(memories)
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[MEMORY] Error getting memories: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/memory")
async def create_memory(
    request: dict,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None
):
    """Create a new memory."""
    if not AUTH_AVAILABLE or not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    try:
        from db_models import MemoryType
        
        memory_type = MemoryType(request.get("memory_type", "fact"))
        memory_key = request.get("memory_key")
        memory_value = request.get("memory_value", {})
        importance_score = request.get("importance_score", 50)
        expires_at = request.get("expires_at")
        
        if not memory_key:
            raise HTTPException(status_code=400, detail="memory_key is required")
        
        if expires_at:
            expires_at = datetime.fromisoformat(expires_at.replace('Z', '+00:00'))
        
        async for db in get_db():
            memory = await user_query_history_service.save_memory(
                db=db,
                user_id=current_user.id,
                memory_type=memory_type,
                memory_key=memory_key,
                memory_value=memory_value,
                importance_score=importance_score,
                expires_at=expires_at
            )
            await db.commit()
            
            return {
                "success": True,
                "memory": {
                    "id": memory.id,
                    "type": memory.memory_type.value,
                    "key": memory.memory_key,
                    "value": memory.memory_value,
                    "importance": memory.importance_score
                }
            }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[MEMORY] Error creating memory: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/memory/{memory_id}")
async def delete_memory(
    memory_id: int,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None
):
    """Delete a memory."""
    if not AUTH_AVAILABLE or not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    try:
        async for db in get_db():
            deleted = await user_query_history_service.delete_memory(
                db=db,
                user_id=current_user.id,
                memory_id=memory_id
            )
            await db.commit()
            
            if not deleted:
                raise HTTPException(status_code=404, detail="Memory not found")
            
            return {"success": True, "message": "Memory deleted"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[MEMORY] Error deleting memory: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/context")
async def get_query_context(
    session_id: Optional[str] = None,
    current_user: User = Depends(require_chatbot_access()) if AUTH_AVAILABLE else None
):
    """Get context for AI agent (recent queries + memories)."""
    if not AUTH_AVAILABLE or not current_user:
        raise HTTPException(status_code=401, detail="Authentication required")
    
    try:
        async for db in get_db():
            context = await user_query_history_service.get_context_for_query(
                db=db,
                user_id=current_user.id,
                session_id=session_id
            )
            
            return {
                "success": True,
                "context": context
            }
    except Exception as e:
        logger.error(f"[CONTEXT] Error getting context: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

